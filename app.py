
from flask import Flask, request, render_template, jsonify , flash, redirect, url_for, session, logging ,send_file
from flask import redirect,url_for
import numpy as np
import pandas as pd
from keras.models import load_model
from datetime import datetime
import joblib

# pdf 
from fpdf import FPDF  # PDF generation
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx2pdf import convert
import os


import tensorflow as tf
import pickle


from flask_mysqldb import MySQL
from wtforms import Form, StringField, TextAreaField, PasswordField, validators
from passlib.hash import sha256_crypt
from functools import wraps
import MySQLdb.cursors

# fill pdf 
# from fillpdf import fillpdfs
# form_fields = list(fillpdfs.get_form_fields('Final Maintenance Report_Template.pdf').keys())

# print(form_fields)

app = Flask(__name__)




# import model
# model = load_model('Model7/ffnn3.keras')
# scaler_X = pickle.load(open('Model7/scaler_X.pkl', 'rb'))
# scaler_y = pickle.load(open('Model7/scaler_y.pkl', 'rb'))




# model / Schema
class RegisterForm(Form):
    name = StringField('Name', [validators.Length(min=1, max=50)])
    username = StringField('Username', [validators.Length(min=4, max=25)])
    email = StringField('Email', [validators.Length(min=6, max=50)])
    password = PasswordField('Password', [
        validators.DataRequired(),
        validators.EqualTo('confirm', message='Passwords do not match')
    ])
    confirm = PasswordField('Confirm Password')


@app.route('/login', methods=['GET', 'POST'])
def login():
    msg = ''
    if request.method == 'POST' and 'username' in request.form and 'password' in request.form:
        username = request.form['username']
        password = request.form['password']
        
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT * FROM accounts WHERE username = %s', (username,))
        account = cursor.fetchone()
        
        if account and bcrypt.checkpw(password.encode('utf-8'), account['password'].encode('utf-8')):
            session['loggedin'] = True
            session['id'] = account['id']
            session['username'] = account['username']
            msg = 'Logged in successfully!'
            return redirect(url_for('home'))  # Redirect to a protected route upon successful login
        else:
            msg = 'Incorrect username / password!'
    
    return render_template('login.html', msg=msg)



def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('loggedin'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


@app.route('/logout')
@login_required
def logout():
    session.pop('loggedin', False)
    session.pop('id', None)
    session.pop('username', None)
    return redirect(url_for('login'))



@app.route('/register', methods=['GET', 'POST'])
def register():
    msg = ''
    if request.method == 'POST' and 'username' in request.form and 'password' in request.form and 'email' in request.form:
        username = request.form['username']
        password = request.form['password']
        email = request.form['email']
        phone_number = request.form.get('phone_number', '')  # Optional field handling
        
        if phone_number and not phone_number.startswith('+91'):
            phone_number = '+91' + phone_number
        # Hash the password
        hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
        
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT * FROM accounts WHERE username = %s', (username,))
        account = cursor.fetchone()

        
        if account:
            msg = 'Account already exists!'
        elif not re.match(r'[^@]+@[^@]+\.[^@]+', email):
            msg = 'Invalid email address!'
        elif not re.match(r'[A-Za-z0-9]+', username):
            msg = 'Username must contain only characters and numbers!'
        elif not username or not password or not email:
            msg = 'Please fill out the form!'
        else:
            cursor.execute(
                'INSERT INTO accounts (username, password, email, phone_number) VALUES (%s, %s, %s, %s)',
                (username, hashed_password, email, phone_number)
            )
            mysql.connection.commit()
            msg = 'You have successfully registered!'
    elif request.method == 'POST':
        msg = 'Please fill out the form!'
    
    return render_template('login.html', msg=msg)

@app.route('/', methods=['GET'])
# @login_required
def home():
    return render_template('home.html')

@app.route('/home2', methods=['GET'])
# @login_required
def home2():
    return render_template('home2.html')

@app.route('/form', methods=['GET'])
# @login_required
def form():
    return render_template('form.html')


@app.route("/info", methods=['GET', 'POST'])
# @login_required
def info():
    return render_template('diseases.html')   


# @app.route('/predict', methods=['POST'])
# @login_required
# def predict():
#     try:
#         # Get input values from the form
#         frequency = float(request.form.get('frequency'))
#         amplitude = float(request.form.get('amplitude'))
#         temperature = float(request.form.get('temperature'))
#         operating_hours = float(request.form.get('OperatingHours'))
#         print(frequency, amplitude, temperature, operating_hours)

#         # Convert input to a DataFrame with column names
#         input_data = pd.DataFrame([[amplitude, frequency, temperature, operating_hours]], 
#                                   columns=['Amplitude', 'Frequency', 'Temperature', 'OperatingHours'])

#         # Scale the input data using the loaded scaler
#         input_data_scaled = scaler_X.transform(input_data)

#         # Make prediction and inverse transform the result
#         prediction_scaled = model.predict(input_data_scaled)
#         predictions = scaler_y.inverse_transform(prediction_scaled)

#         # Convert predictions to native Python types
#         predicted_mass = float(predictions[0, 0])
#         predicted_lifespan = float(predictions[0, 1])
#         predicted_unbalance_force = float(predictions[0, 2])

#         # Define severity bins, labels, and map
#         bins = [0, 5, 10, 25, 40, 55, 65, np.inf]
#         labels = ['Negligible', 'Minor', 'Moderate', 'Significant', 'Serious', 'Severe', 'Critical']
#         severity_map = {'Negligible': 0, 'Minor': 1, 'Moderate': 2, 'Significant': 3, 'Serious': 4, 'Severe': 5, 'Critical': 6}

#         # Determine severity based on predicted mass
#         severity = pd.cut([predicted_mass], bins=bins, labels=labels, include_lowest=True).to_list()[0]
#         severity_numerical = severity_map[severity]

#         # Prepare data for rendering the template
#         result = {
#             'predicted_mass': predicted_mass,
#             'predicted_unbalance_force': predicted_unbalance_force,
#             'predicted_lifespan': predicted_lifespan,
#             'severity_name': severity,
#             'severity_numerical': severity_numerical
#         }

#         print(result)

#         user_id = session.get('id')
#         print(user_id)
#         # Check if the severity_numerical is 6 (Critical)
#         if severity_numerical == 6:
#             # Send email to the user
#             query = "SELECT phone_number, email FROM accounts WHERE id = %s"
#             cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
#             cursor.execute(query, (user_id,))
#             user = cursor.fetchone()

#             if user:
#                phone_number = user.get('phone_number')
#                email = user.get('email')
#                print(phone_number)  
#                print(email)
                
#             msg = Message(
#                 "Machine Maintenance Required",
#                 sender="derugaud@gmail.com",
#                 recipients=[email]
#             )
#             msg.body = "The machine needs maintenance immediately due to critical mass prediction."
            
#             # send_sms(phone_number)
#             # mail.send(msg)

#         return render_template('result.html', **result)

#     except Exception as e:
#         # Handle any exceptions that occur during processing
#         return f"Error during prediction: {str(e)}"

# rf_model = load_model('modelRandom/random_forest_model.pkl')
# rf_scaler_X = pickle.load(open('modelRandom/scaler_X.pkl', 'rb'))
# rf_scaler_y = pickle.load(open('modelRandom/scaler_y.pkl', 'rb'))

rf_model = joblib.load('modelRandom/random_forest_model.pkl')

# Load the scalers for X and y
# rf_scaler_X = joblib.load('modelRandom/scaler_X.pkl')
# rf_scaler_y = joblib.load('modelRandom/scaler_y.pkl')


@app.route('/predict', methods=['POST'])
def predict():

    print(request.form)
    print("pramod*********************")
    try:
        # Retrieve form inputs
        airflow_rate = float(request.form['Airflow_Rate'])
        fuel_flowrate = float(request.form['Fuel_Flowrate'])
        air_fuel_mixture = float(request.form['Air_Fuel_Mixture'])
        pressure_sensor_data = float(request.form['Pressure_Sensor_Data'])
        sound_intensity = float(request.form['Sound_Intensity'])
        burner_position = request.form['Burner_Position']

        # Encode burner position
        burner_position_L2 = 1 if burner_position == "L/2" else 0
        burner_position_L4 = 1 if burner_position == "L/4" else 0

        # Create an array for the input features
        input_data = np.array([[airflow_rate, fuel_flowrate, air_fuel_mixture,
                                pressure_sensor_data, sound_intensity,
                                burner_position_L2, burner_position_L4]])

        # Make prediction using the trained Random Forest model without scaling
        predicted_output = rf_model.predict(input_data)
        print(predicted_output[0])
        # Return the predicted result as a rendered HTML page
        return render_template(
    "result.html",
    predicted_value=round(predicted_output[0], 2),
    airflow_rate=airflow_rate,
    fuel_flowrate=fuel_flowrate,
    air_fuel_mixture=air_fuel_mixture,
    pressure_sensor_data=pressure_sensor_data,
    sound_intensity=sound_intensity,
    burner_position=burner_position
)


    except Exception as e:
        return jsonify({'error': str(e)})


def generate_prediction_pdf(result):
    # Define the path to save the document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = f"/mnt/data/Predictive_Maintenance_Report_{timestamp}.docx"
    pdf_path = docx_path.replace(".docx", ".pdf")

    # Create a new document
    doc = Document()
    doc.add_heading('Fault Report: Predictive Maintenance for Rotating Machine', 0)

    # Machine Information
    doc.add_heading('Machine Information', level=1)
    doc.add_paragraph('Machine Name: Rotating Machine')
    doc.add_paragraph('Machine Type: Shaft-based Rotating Machine')
    doc.add_paragraph('Machine No.: RM-001')

    # Data Summary in a table format
    doc.add_heading('Data Summary', level=1)
    data_summary_table = doc.add_table(rows=1, cols=2)
    hdr_cells = data_summary_table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Value'

    # Adding borders to the table
    tbl = data_summary_table._tbl
    tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>' % nsdecls('w'))
    tbl.tblPr.append(tblBorders)

    # Adding data summary values
    data_summary = [
        ('Predicted Mass', f'{result["predicted_mass"]} grams'),
        ('Predicted Unbalance Force', f'{result["predicted_unbalance_force"]} N'),
        ('Predicted Lifespan', f'{result["predicted_lifespan"]} hours'),
        ('Severity Index', f'{result["severity_numerical"]} ({result["severity_name"]})')
    ]
    for item, value in data_summary:
        row_cells = data_summary_table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = value

    # Fault Occurrence
    doc.add_heading('Fault Occurrence', level=1)
    doc.add_paragraph('Fault Type: Shaft Unbalance')
    doc.add_paragraph('Time of Fault: 2024-09-07 10:30 AM')
    doc.add_paragraph('Date of Fault: 2024-09-07')

    # Severity Mapping with color representation
    doc.add_heading('Severity Mapping', level=1)
    severity_table = doc.add_table(rows=1, cols=3)
    hdr_cells = severity_table.rows[0].cells
    hdr_cells[0].text = 'Predicted Mass (g)'
    hdr_cells[1].text = 'Severity'
    hdr_cells[2].text = 'Meaning'

    severity_data = [
        ('0 - 4', '0', 'Negligible', RGBColor(40, 167, 69)),           # Green
        ('5 - 9', '1', 'Minor', RGBColor(111, 187, 111)),              # Light Green
        ('10 - 24', '2', 'Moderate', RGBColor(253, 216, 53)),          # Yellow
        ('25 - 39', '3', 'Significant', RGBColor(242, 158, 36)),       # Orange
        ('40 - 54', '4', 'Serious', RGBColor(243, 115, 33)),           # Dark Orange
        ('55 - 65', '5', 'Severe', RGBColor(240, 58, 23)),             # Red-Orange
        ('Above 65', '6', 'Critical', RGBColor(220, 53, 69))           # Red
    ]

    # Adding severity data with color-coded Meaning descriptions
    for mass, severity, meaning, color in severity_data:
        row_cells = severity_table.add_row().cells
        row_cells[0].text = mass
        row_cells[1].text = severity
        meaning_run = row_cells[2].paragraphs[0].add_run(meaning)
        meaning_run.font.color.rgb = color

    # Adding borders to the severity table
    tbl = severity_table._tbl
    tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>' % nsdecls('w'))
    tbl.tblPr.append(tblBorders)

    # Recommendations and Repair Suggestions
    doc.add_heading('Recommendations and Repair Suggestions', level=1)
    doc.add_paragraph('Recommendation: Reduce operating load and check alignment of the shaft. '
                      'Perform dynamic balancing to reduce unbalance force.')
    doc.add_paragraph('Repair Suggestion: Add or remove mass from the shaft to counteract the imbalance. '
                      'Regularly check the accelerometer readings and monitor temperature closely.')

    # Maintenance Schedule Based on Severity
    doc.add_heading('Maintenance Schedule Based on Severity', level=1)
    maintenance_points = [
        'Severity 0-1 (Negligible/Minor): Perform routine visual inspections every 3 months.',
        'Severity 2 (Moderate): Conduct detailed inspections every 2 months and check unbalance force.',
        'Severity 3 (Significant): Perform dynamic balancing every month.',
        'Severity 4 (Serious): Immediate inspection required. Dynamic balancing and shaft realignment within the week.',
        'Severity 5 (Severe): Immediate action required. Halt operation and replace the shaft if needed.',
        'Severity 6 (Critical): Immediate shutdown of the machine and full maintenance, including rotor replacement.'
    ]
    for point in maintenance_points:
        doc.add_paragraph(f'- {point}', style='ListBullet')

    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(
        "The rotating machine exhibited unbalance due to uneven mass distribution on the shaft. "
        "Severity levels indicate potential issues ranging from minor to critical. Regular maintenance and balancing of "
        "the rotor are essential to prolong the machine's lifespan and ensure smooth operation. Immediate corrective measures "
        "should be taken to avoid further damage and minimize downtime."
    )
    
    # Save the document
    doc.save(docx_path)

    # Convert DOCX to PDF (Ensure convert function is correctly defined/imported)
    convert(docx_path, pdf_path)
    print(docx_path)
    
    
    try:
        bucket_name = 'hackwin'
        s3_file_name = f'reports/{os.path.basename(pdf_path)}'
        
        s3.upload_file(pdf_path, bucket_name, s3_file_name)
        
        # Generate a public URL for the file
        s3_url = f"https://{bucket_name}.s3.amazonaws.com/{s3_file_name}"
        
        # Remove the local files after uploading
        os.remove(docx_path)
        os.remove(pdf_path)
        
        return s3_url
    
    except FileNotFoundError:
        return "The file was not found"
    
    except NoCredentialsError:
        return "Credentials not available"

    

    except Exception as e:
        print(f"Error during prediction: {e}")
        return f"Error occurred: {e}"





def fill_pdf(result):
    # Define the data dictionary to fill the PDF fields
    data_dict = {
        'MI': "",  # Example value for Machine ID
        'SN': "",  # Example value for Serial Number
        'MN': "",  # Example value for Machine Name
        'Manu': "",  # Example value for Manufacturer
        'date': datetime.now().strftime('%Y-%m-%d'),  # Current date
        'Depart': "",  # Example value for Department
        'freq': result.get("frequency", ""),  # Frequency
        'amp': result.get("amplitude", ""),  # Amplitude
        'mass': result.get("predicted_mass", ""),  # Predicted Mass
        'severity': f"{result.get('severity_name', '')} {result.get('severity_numerical', '')}",  # Severity
        'unbalance': result.get("predicted_unbalance_force", ""),  # Predicted Unbalance Force
        'temp': result.get("temperature", ""),  # Temperature
        'operating': result.get("operating_hours", ""),  # Operating Hours
        'lifespan': result.get("predicted_lifespan", ""),  # Predicted Lifespan
        'fault type': "",  # Fault Type
        'Date of occurance': "",  # Date of Occurrence
        'Time of occurance': "",  # Time of Occurrence
        '0-1': "",  # Fault codes
        '2': "",
        '3': "",
        '4': "",
        '5': "",
        '6': "",
        'summary': "",  # Summary
        'remark': "",  # Remarks
        'maintenence Engineer': "",  # Maintenance Engineer
        'manager': "",  # Manager
        'operator': "",  # Operator
        'inspector': "",  # Inspector
        'd &  t 1': "",  # Dates and Times
        'd & t 2': "",
        'd & t 3': "",
        'd & t 4': ""
    }

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    name = f"Predictive_Maintenance_Report_{timestamp}.pdf"

    # Fill the PDF with data from the dictionary and create a new file
    fillpdfs.write_fillable_pdf('Final Maintenance Report_Template.pdf', name, data_dict, flatten=True)

    # Define the S3 client
    
    try:
        # Upload the file to S3
        bucket_name = 'hackwin'
        s3_file_name = f'reports/{os.path.basename(name)}'
        s3.upload_file(name, bucket_name, s3_file_name)
        
        # Generate a public URL for the file
        s3_url = f"https://{bucket_name}.s3.amazonaws.com/{s3_file_name}"
        
        # Remove the local file after uploading
        os.remove(name)
        
        return s3_url
    except FileNotFoundError:
        return "The file was not found"
    except NoCredentialsError:
        return "Credentials not available"
    except Exception as e:
        # Ensure the local file is deleted even if an error occurs
        if os.path.exists(name):
            os.remove(name)
        return f"An error occurred: {str(e)}"



@app.route('/download_pdf')
def download_pdf():
    pdf_file_path = request.args.get('pdf_file_path')
    if os.path.exists(pdf_file_path):
        return send_file(pdf_file_path, as_attachment=True)
    return "File not found."



def send_sms(phone_number):
    responseData = sms.send_message(
      {
        "from": "DAYNA4MITE",
        "to": phone_number,
        "text": "The machine needs maintenance immediately due to critical mass prediction.",
      }
    )

    if responseData["messages"][0]["status"] == "0":
      print("Message sent successfully.")
    else:
      print(f"Message failed with error: {responseData['messages'][0]['error-text']}")

if __name__ == '__main__':
    app.secret_key='secret123'
    app.run(debug=True,port=5000)